from PyPDF2 import PdfReader
from openpyxl import load_workbook
import os
from pathlib import Path
from docx import Document
from openpyxl import Workbook
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import zipfile
import rarfile
import random
import shutil
import pandas as pd
import rarfile
import xlwt

# Создаем папку documents, если её нет
Path("documents").mkdir(exist_ok=True)

def generate_text():
    """Генерирует случайный текст для документов"""
    texts = [
        
        "Это тестовый документ для проверки работы поискового краулера.",
        "В этом файле содержится важная информация о проекте.",
        "Данные для анализа представлены в табличной форме.",
        "Полнотекстовый поиск должен находить этот контент.",
        "Python скрипт создал этот файл автоматически."
    ]
    return random.choice(texts)

def create_docx(filename):
    """Создает DOCX файл"""
    doc = Document()
    doc.add_paragraph(generate_text())
    doc.add_paragraph("Вторая строка с данными.")
    doc.save(filename)

def create_doc(filename):
    """Создает DOC файл (через сохранение как .doc)"""
    doc = Document()
    doc.add_paragraph("Это устаревший DOC формат.")
    doc.add_paragraph(generate_text())
    doc.save(filename)

def create_xlsx(filename):
    """Создает XLSX файл"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Тестовые данные"
    ws.append(["ID", "Name", "Value"])
    for i in range(1, 4):
        ws.append([i, f"Item {i}", random.randint(10, 100)])
    wb.save(filename)

def create_xls(filename):
    """Создает XLS файл (старый формат)"""
    # Для .xls нужно использовать библиотеку xlwt
    try:
        
        wb = xlwt.Workbook()
        ws = wb.add_sheet("Data")
        ws.write(0, 0, "Это старый XLS формат")
        ws.write(1, 0, generate_text())
        wb.save(filename)
    except ImportError:
        print("Для создания .xls установите xlwt: pip install xlwt")

def create_pdf(filename):
    """Создает PDF файл"""
    c = canvas.Canvas(filename, pagesize=letter)
    c.drawString(100, 750, generate_text())
    c.drawString(100, 730, "PDF документ для тестирования.")
    c.save()

def create_archive(archive_name, files_to_archive, archive_type="zip"):
    """Создает архив с файлами"""
    if archive_type == "zip":
        with zipfile.ZipFile(archive_name, 'w') as zipf:
            for file in files_to_archive:
                zipf.write(file, os.path.basename(file))
    # elif archive_type == "rar":
    #     with rarfile.RarFile(archive_name, 'w') as rarf:
    #         for file in files_to_archive:
    #             rarf.write(file, os.path.basename(file))
    # elif archive_type == "7z":
    #     with py7zr.SevenZipFile(archive_name, 'w') as szf:
    #         for file in files_to_archive:
    #             szf.write(file, os.path.basename(file)) ВОзникли проблемы с утановкой библиотеки 

def main():
    # Генерируем основные файлы
    files_created = []
    
    office_files = [
        ("docx_file.docx", create_docx),
        ("doc_file.doc", create_doc),
        ("xlsx_file.xlsx", create_xlsx),
        ("xls_file.xls", create_xls),
        ("pdf_file.pdf", create_pdf),
    ]
    
    for filename, creator in office_files:
        path = os.path.join("documents", filename)
        creator(path)
        files_created.append(path)
        print(f"Created: {path}")
    
    # Создаем архив с некоторыми из этих файлов
    archive_files = random.sample(files_created, 3)
    
    archive_types = ["zip"]
    for arch_type in archive_types:
        archive_name = os.path.join("documents", f"archive_{arch_type}.{arch_type}")
        create_archive(archive_name, archive_files, arch_type)
        print(f"Created archive: {archive_name}")
    
    # Создаем вложенные архивы (архив в архиве)
    nested_archive_path = os.path.join("documents", "nested_archive.zip")
    with zipfile.ZipFile(nested_archive_path, 'w') as zipf:
        for file in random.sample(files_created, 2):
            zipf.write(file, os.path.basename(file))
        # Добавляем существующий архив внутрь
        zipf.write(
            os.path.join("documents", "archive_zip.zip"),
            "inner_archive.zip"
        )
    print(f"Created nested archive: {nested_archive_path}")

if __name__ == "__main__":
    main()